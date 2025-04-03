import requests
import json
from datetime import datetime, timezone
from docx import Document
from pathlib import Path
import os

#Converts date to strandard and returns true if current date is not past due date
def assignmentPastDue(date_str):
    given_time = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
    current_time = datetime.now(timezone.utc)
    if current_time < given_time:
        return(True)
    else:
        return(False)
    
#Helper function to parse and shorten name of assignment
def shortenName(assignmentTitle):
    input_string = assignmentTitle
    parts = input_string.split()
    desired_text = " ".join(parts[3:]).split("(")[0].strip()
    return(desired_text)

#Helper function to prettify date
def makeDateNice(inputDate):
    date_string = inputDate
    date_obj = datetime.strptime(date_string, "%Y-%m-%dT%H:%M:%SZ")
    formatted_date = date_obj.strftime("%m/%d/%Y")
    return(formatted_date)

#Lab Template written into Doc
def writeDoc(shortenedName,dueDate,courseName, studentName):
    subfolder = Path("output_docs")  # Name of the subfolder
    subfolder.mkdir(parents=True, exist_ok=True)  # Create subfolder if it doesn't exist
    fName=shortenedName+".docx"
    file_path = subfolder / fName  # Full path to the file
    doc = Document()
    doc.add_paragraph("Name: " + studentName)
    doc.add_paragraph("Course Name: " + courseName)
    doc.add_paragraph("Date: " + dueDate)
    doc.add_paragraph("Assignment: " + shortenedName +"\n\n")
    doc.add_paragraph("Purpose: \n\n")
    doc.add_paragraph("Environment: \n\n")
    doc.add_paragraph("Lab Diagram: \n\n")
    doc.add_paragraph("Problems Encountered: \n\n")
    doc.add_paragraph("Lessons Learned and Discussion: \n\n")
    doc.add_paragraph("Lab Screenshots: \n\n")
    doc.add_paragraph("Auto-Grader Screenshots: \n\n")
    doc.save(file_path)
def most_frequent_number(arr):
    counts = {}
    
    for num in arr:
        counts[num] = counts.get(num, 0) + 1  # Count occurrences

    most_common = max(counts, key=counts.get)  # Find the number with the highest count
    print("Most common number: "+ str(most_common) + " Count of most common out of 7: " + str(counts[most_common]))
    return most_common, counts[most_common]
def most_frequent_number(arr):
    counts = {}
    
    for num in arr:
        counts[num] = counts.get(num, 0) + 1  # Count occurrences

    most_common = max(counts, key=counts.get)  # Find the number with the highest count
    print("Most common number: "+ str(most_common) + " Count of most common out of 7: " + str(counts[most_common]))
    return most_common
#Takes in Class code and returns canvas api link with assignment_group_id and class name in a tuple
def getClassInfo(bToken):
    userClassName=input("Enter class code:")
    r=requests.get("https://canvas.instructure.com/api/v1/courses?per_page=100", headers={"Authorization":"Bearer " + bToken})
    class_json=r.content
    class_dict=json.loads(class_json)
    #print(class_dict)
    for class1 in class_dict:
            #print(class1.get('end_at'))
            #print(class1.get('course_code'))
            if((class1.get('course_code'))==userClassName):
                print("Sucessfully matched course code")
                print(class1.get('course_code'))
                #print(class1.get('id'))
                userClassID=str(class1.get('id'))
                d=requests.get("https://canvas.instructure.com/api/v1/courses/" + userClassID + "/assignments?search_term=Lab&per_page=7", headers={"Authorization":"Bearer " + bToken})
                assignmentsSearch_json=d.content
                assignmentsSearch_dict=json.loads(assignmentsSearch_json)
                assignments_arr=[]
                for assignmentCurrent in assignmentsSearch_dict:
                    print(assignmentCurrent.get('assignment_group_id'))
                    assignments_arr.append(assignmentCurrent.get('assignment_group_id'))
                mostFreq=most_frequent_number(assignments_arr)
                print("https://canvas.instructure.com/api/v1/courses/" + userClassID + "/assignments?assignment_group_id=" + str(mostFreq))
                return ("https://canvas.instructure.com/api/v1/courses/" + userClassID + "/assignments?assignment_group_id=" + str(mostFreq), userClassName)
def main():
    file_path = Path("doNotDeleteBearTokenFile.txt")
    if file_path.exists():
        f = open("doNotDeleteBearTokenFile.txt", "r")
        tkn=f.read()
        f.close()
        bearerToken=tkn
    else:
        bearerToken=input("Enter Canvas Token (See README on how to get code): ")
        f = open("doNotDeleteBearTokenFile.txt", "w")
        f.write(bearerToken)
        f.close()
        f = open("doNotDeleteBearTokenFile.txt", "r")
        tkn=f.read()
        f.close()
        bearerToken=tkn
    
    if(len(bearerToken)!=69):
        print("Token format incorrect")
        os.remove("doNotDeleteBearTokenFile.txt")
        exit()
    userName=input("Enter full name, first name last name: ")
    classesTupe=getClassInfo(bearerToken)
    r=requests.get(classesTupe[0], headers={"Authorization":"Bearer " + bearerToken})
    assignments_json=r.content
    assignments_dict=json.loads(assignments_json)
    print(type(assignments_dict))
    for assignment in assignments_dict:
        if(assignmentPastDue(assignment.get('due_at'))):
            print(assignment.get('name'))
            assignmentName=assignment.get('name')
            print(assignment.get('due_at'))
            assignmentDueDate=assignment.get('due_at')
            print(shortenName(assignmentName))
            writeDoc(shortenName(assignmentName),makeDateNice(assignmentDueDate), classesTupe[1], userName)

if __name__ == '__main__':
    main()










